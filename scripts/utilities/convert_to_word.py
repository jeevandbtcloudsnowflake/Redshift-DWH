#!/usr/bin/env python3
"""
Convert Markdown to Word Document
Converts the VPC usage conversation details to a professional Word document
"""

import os
import sys
from pathlib import Path

def install_requirements():
    """Install required packages for Word document generation"""
    packages = [
        'python-docx',
        'markdown',
        'beautifulsoup4'
    ]
    
    for package in packages:
        try:
            if package == 'python-docx':
                import docx
            elif package == 'markdown':
                import markdown
            elif package == 'beautifulsoup4':
                import bs4
            print(f"✅ {package} already installed")
        except ImportError:
            print(f"📦 Installing {package}...")
            import subprocess
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])

def create_word_document():
    """Create a professional Word document from the Markdown content"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import markdown
        from bs4 import BeautifulSoup
        import re
        
        # Read the markdown file
        md_file = Path('docs/VPC_Usage_Conversation_Details.md')
        if not md_file.exists():
            print(f"❌ Markdown file not found: {md_file}")
            return False
        
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('VPC Usage in AWS Data Warehouse', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_heading('Complete Conversation Analysis & Technical Documentation', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document info
        doc.add_paragraph()
        info_para = doc.add_paragraph()
        info_para.add_run('Document Created: ').bold = True
        info_para.add_run('December 2024\n')
        info_para.add_run('Project: ').bold = True
        info_para.add_run('E-commerce Data Warehouse\n')
        info_para.add_run('Topic: ').bold = True
        info_para.add_run('VPC with Private Subnets Usage Analysis')
        
        doc.add_page_break()
        
        # Process markdown content
        lines = md_content.split('\n')
        current_section = None
        in_code_block = False
        code_content = []
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines and document header
            if not line or line.startswith('**Document Created**'):
                continue
            
            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    if code_content:
                        code_para = doc.add_paragraph()
                        code_run = code_para.add_run('\n'.join(code_content))
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(9)
                        # Add background color (light gray)
                        code_para.style = 'Intense Quote'
                    code_content = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(line)
                continue
            
            # Handle headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('#').strip()
                
                # Skip table of contents and main title
                if 'Table of Contents' in heading_text or heading_text == 'VPC Usage in AWS Data Warehouse - Complete Conversation Analysis':
                    continue
                
                if level <= 3:
                    doc.add_heading(heading_text, level)
                else:
                    # For deeper levels, use bold paragraph
                    para = doc.add_paragraph()
                    run = para.add_run(heading_text)
                    run.bold = True
                    run.font.size = Pt(12)
                continue
            
            # Handle bullet points
            if line.startswith('- ') or line.startswith('* '):
                bullet_text = line[2:].strip()
                # Clean up markdown formatting
                bullet_text = re.sub(r'\*\*(.*?)\*\*', r'\1', bullet_text)  # Remove bold
                bullet_text = re.sub(r'`(.*?)`', r'\1', bullet_text)  # Remove code formatting
                doc.add_paragraph(bullet_text, style='List Bullet')
                continue
            
            # Handle numbered lists
            if re.match(r'^\d+\.', line):
                list_text = re.sub(r'^\d+\.\s*', '', line)
                list_text = re.sub(r'\*\*(.*?)\*\*', r'\1', list_text)
                doc.add_paragraph(list_text, style='List Number')
                continue
            
            # Handle regular paragraphs
            if line and not line.startswith('---'):
                # Clean up markdown formatting
                clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Bold
                clean_line = re.sub(r'`(.*?)`', r'\1', clean_line)  # Code
                clean_line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean_line)  # Links
                
                # Add paragraph
                para = doc.add_paragraph(clean_line)
                
                # Apply formatting for special content
                if line.startswith('**') and line.endswith('**'):
                    para.runs[0].bold = True
                elif '✅' in line or '❌' in line:
                    para.style = 'List Bullet'
        
        # Add conclusion section
        doc.add_page_break()
        doc.add_heading('Document Summary', 1)
        
        summary_points = [
            "VPC with private subnets is mandatory for Amazon Redshift deployment",
            "This implementation follows industry security best practices",
            "The architecture is cost-effective (VPC components are free)",
            "This pattern is used universally across enterprise environments",
            "The implementation is production-ready and compliant with regulations"
        ]
        
        for point in summary_points:
            doc.add_paragraph(f"• {point}", style='List Bullet')
        
        # Save the document
        output_file = Path('docs/VPC_Usage_Analysis_Complete.docx')
        doc.save(str(output_file))
        
        print(f"✅ Word document created: {output_file}")
        print(f"📄 Document contains comprehensive VPC usage analysis")
        print(f"📊 Includes technical details, cost analysis, and learning resources")
        
        return True
        
    except Exception as e:
        print(f"❌ Error creating Word document: {str(e)}")
        return False

def create_simple_word_document():
    """Create a simple Word document without complex formatting"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        # Read the markdown file
        md_file = Path('docs/VPC_Usage_Conversation_Details.md')
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create simple document
        doc = Document()
        
        # Add title
        doc.add_heading('VPC Usage in AWS Data Warehouse - Complete Analysis', 0)
        
        # Add content as paragraphs
        sections = content.split('\n## ')
        
        for section in sections:
            if not section.strip():
                continue
                
            lines = section.split('\n')
            if lines:
                # Add section heading
                if not section.startswith('#'):
                    doc.add_heading(lines[0], 1)
                    lines = lines[1:]
                
                # Add content
                current_para = ""
                for line in lines:
                    line = line.strip()
                    if line.startswith('###'):
                        if current_para:
                            doc.add_paragraph(current_para)
                            current_para = ""
                        doc.add_heading(line.replace('###', '').strip(), 2)
                    elif line.startswith('**') and line.endswith('**'):
                        if current_para:
                            doc.add_paragraph(current_para)
                            current_para = ""
                        para = doc.add_paragraph()
                        run = para.add_run(line.replace('**', ''))
                        run.bold = True
                    elif line and not line.startswith('```') and not line.startswith('---'):
                        if current_para:
                            current_para += " " + line
                        else:
                            current_para = line
                
                if current_para:
                    doc.add_paragraph(current_para)
        
        # Save document
        output_file = Path('docs/VPC_Usage_Analysis_Simple.docx')
        doc.save(str(output_file))
        
        print(f"✅ Simple Word document created: {output_file}")
        return True
        
    except Exception as e:
        print(f"❌ Error creating simple Word document: {str(e)}")
        return False

def main():
    """Main function to create Word document"""
    print("📄 Converting VPC Usage Analysis to Word Document")
    print("=" * 60)
    
    # Install requirements
    print("🔧 Installing required packages...")
    install_requirements()
    
    # Try to create comprehensive document first
    print("\n📝 Creating comprehensive Word document...")
    if create_word_document():
        print("✅ Comprehensive document created successfully!")
    else:
        print("⚠️ Comprehensive document failed, trying simple version...")
        if create_simple_word_document():
            print("✅ Simple document created successfully!")
        else:
            print("❌ Both document creation methods failed")
            return False
    
    print("\n🎯 Word Document Creation Complete!")
    print("📁 Output files:")
    
    # List created files
    doc_files = list(Path('docs').glob('*.docx'))
    for doc_file in doc_files:
        if 'VPC_Usage' in doc_file.name:
            print(f"   • {doc_file}")
    
    print("\n💡 You can now:")
    print("   1. Open the Word document in Microsoft Word")
    print("   2. Edit formatting as needed")
    print("   3. Share with team members")
    print("   4. Print or convert to PDF")
    
    return True

if __name__ == "__main__":
    main()
